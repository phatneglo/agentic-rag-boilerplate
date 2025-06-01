"""
Simplified document converter worker.
Processes document conversion jobs from the Redis queue using basic conversion methods.
This is a fallback version that works without Marker library dependencies.
"""
import asyncio
import os
import sys
from pathlib import Path
from typing import Any, Dict
import tempfile
import shutil
from datetime import datetime

# Add the parent directory to the Python path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from bullmq import Worker
import redis.asyncio as redis

from app.core.config import settings
from app.core.logging_config import configure_logging, get_logger, log_job_event
from app.utils.exceptions import DocumentConversionError, FileProcessingError

# Configure logging
configure_logging()
logger = get_logger(__name__)


class SimpleDocumentConverterWorker:
    """Worker for processing document conversion jobs using basic methods."""
    
    def __init__(self):
        self.worker = None
        self.redis_connection = None
    
    async def setup(self):
        """Setup Redis connection and worker."""
        try:
            # Create Redis connection for health checks
            self.redis_connection = redis.Redis(
                host=settings.redis_host,
                port=settings.redis_port,
                password=settings.redis_password,
                db=settings.redis_db,
                decode_responses=True,
            )
            
            # Test connection
            await self.redis_connection.ping()
            logger.info("Redis connection established for simple document converter worker")
            
            # Create worker - BullMQ Python API
            self.worker = Worker(
                settings.queue_names["document_converter"],
                self.process_job,
            )
            
            logger.info(
                "Simple document converter worker initialized",
                queue_name=settings.queue_names["document_converter"]
            )
            
        except Exception as e:
            logger.error("Failed to setup simple document converter worker", error=str(e))
            raise
    
    async def process_job(self, job) -> Dict[str, Any]:
        """
        Process a document conversion job using basic conversion methods.
        
        Args:
            job: BullMQ job object
            
        Returns:
            Dict[str, Any]: Job result
            
        Raises:
            DocumentConversionError: If conversion fails
        """
        job_id = job.id
        job_data = job.data
        
        try:
            logger.info(
                "Processing simple document conversion job",
                job_id=job_id,
                document_id=job_data.get("document_id"),
                source_path=job_data.get("source_path")
            )
            
            log_job_event(
                job_id=job_id,
                queue_name=settings.queue_names["document_converter"],
                event_type="started",
                document_id=job_data.get("document_id")
            )
            
            # Extract job data
            document_id = job_data["document_id"]
            source_path = job_data["source_path"]
            output_path = job_data["output_path"]
            conversion_options = job_data.get("conversion_options", {})
            
            # Update job progress
            await job.updateProgress(10)
            
            # Validate source file exists
            if not os.path.exists(source_path):
                raise DocumentConversionError(f"Source file not found: {source_path}")
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            os.makedirs(output_dir, exist_ok=True)
            
            await job.updateProgress(20)
            
            # Convert document based on file type
            source_ext = os.path.splitext(source_path.lower())[1]
            
            if source_ext == '.pdf':
                result = await self._convert_pdf_simple(
                    source_path, output_path, conversion_options
                )
            elif source_ext in ['.docx', '.doc']:
                result = await self._convert_docx_simple(
                    source_path, output_path, conversion_options
                )
            elif source_ext in ['.txt', '.md']:
                result = await self._convert_text_to_markdown(
                    source_path, output_path, conversion_options
                )
            elif source_ext in ['.html', '.htm']:
                result = await self._convert_html_to_markdown(
                    source_path, output_path, conversion_options
                )
            else:
                # Fallback: copy as text and wrap in markdown
                result = await self._fallback_conversion(
                    source_path, output_path, conversion_options
                )
            
            await job.updateProgress(90)
            
            # Prepare result
            job_result = {
                "success": True,
                "document_id": document_id,
                "source_path": source_path,
                "output_path": output_path,
                "conversion_result": result,
                "processed_at": datetime.utcnow().isoformat(),
            }
            
            await job.updateProgress(100)
            
            logger.info(
                "Simple document conversion job completed successfully",
                job_id=job_id,
                document_id=document_id,
                output_path=output_path
            )
            
            log_job_event(
                job_id=job_id,
                queue_name=settings.queue_names["document_converter"],
                event_type="completed",
                document_id=document_id,
                result=job_result
            )
            
            return job_result
            
        except Exception as e:
            logger.error(
                "Simple document conversion job failed",
                job_id=job_id,
                document_id=job_data.get("document_id"),
                error=str(e)
            )
            
            log_job_event(
                job_id=job_id,
                queue_name=settings.queue_names["document_converter"],
                event_type="failed",
                document_id=job_data.get("document_id"),
                error=str(e)
            )
            
            raise DocumentConversionError(f"Simple document conversion failed: {e}")
    
    async def _convert_pdf_simple(
        self,
        source_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Simple PDF conversion using PyPDF2."""
        try:
            import PyPDF2
            
            def convert_pdf():
                with open(source_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = []
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"## Page {page_num}\n\n{page_text}\n")
                    
                    markdown_content = f"# Document: {os.path.basename(source_path)}\n\n" + "\n".join(text_content)
                    
                    with open(output_path, 'w', encoding='utf-8') as md_file:
                        md_file.write(markdown_content)
                    
                    return {
                        "pages_processed": len(pdf_reader.pages),
                        "method": "PyPDF2",
                        "content_length": len(markdown_content)
                    }
            
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, convert_pdf)
            
            logger.info("PDF converted successfully using PyPDF2", result=result)
            return result
            
        except ImportError:
            # Fallback if PyPDF2 not available
            logger.warning("PyPDF2 not available, using fallback conversion")
            return await self._fallback_conversion(source_path, output_path, options)
        except Exception as e:
            logger.error("PDF conversion failed", error=str(e))
            raise DocumentConversionError(f"PDF conversion failed: {e}")
    
    async def _convert_docx_simple(
        self,
        source_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Simple DOCX conversion using python-docx."""
        try:
            from docx import Document
            
            def convert_docx():
                doc = Document(source_path)
                content_parts = [f"# Document: {os.path.basename(source_path)}\n"]
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text + "\n")
                
                markdown_content = "\n".join(content_parts)
                
                with open(output_path, 'w', encoding='utf-8') as md_file:
                    md_file.write(markdown_content)
                
                return {
                    "paragraphs_processed": len([p for p in doc.paragraphs if p.text.strip()]),
                    "method": "python-docx",
                    "content_length": len(markdown_content)
                }
            
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, convert_docx)
            
            logger.info("DOCX converted successfully", result=result)
            return result
            
        except ImportError:
            logger.warning("python-docx not available, using fallback conversion")
            return await self._fallback_conversion(source_path, output_path, options)
        except Exception as e:
            logger.error("DOCX conversion failed", error=str(e))
            raise DocumentConversionError(f"DOCX conversion failed: {e}")
    
    async def _convert_text_to_markdown(
        self,
        source_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Convert text/markdown files."""
        try:
            with open(source_path, 'r', encoding='utf-8') as source_file:
                content = source_file.read()
            
            # If it's already markdown, just copy with header
            if source_path.lower().endswith('.md'):
                markdown_content = content
            else:
                # Convert plain text to markdown
                markdown_content = f"# Document: {os.path.basename(source_path)}\n\n```\n{content}\n```"
            
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write(markdown_content)
            
            result = {
                "method": "text-copy" if source_path.lower().endswith('.md') else "text-to-markdown",
                "content_length": len(markdown_content),
                "original_length": len(content)
            }
            
            logger.info("Text/Markdown conversion completed", result=result)
            return result
            
        except Exception as e:
            logger.error("Text conversion failed", error=str(e))
            raise DocumentConversionError(f"Text conversion failed: {e}")
    
    async def _convert_html_to_markdown(
        self,
        source_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Convert HTML to markdown."""
        try:
            from markdownify import markdownify as md
            
            with open(source_path, 'r', encoding='utf-8') as html_file:
                html_content = html_file.read()
            
            markdown_content = md(html_content)
            
            with open(output_path, 'w', encoding='utf-8') as md_file:
                md_file.write(markdown_content)
            
            result = {
                "method": "markdownify",
                "content_length": len(markdown_content),
                "original_length": len(html_content)
            }
            
            logger.info("HTML converted to Markdown", result=result)
            return result
            
        except ImportError:
            logger.warning("markdownify not available, using fallback")
            return await self._fallback_conversion(source_path, output_path, options)
        except Exception as e:
            logger.error("HTML conversion failed", error=str(e))
            raise DocumentConversionError(f"HTML conversion failed: {e}")
    
    async def _fallback_conversion(
        self,
        source_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Fallback conversion - treat as text and wrap in markdown."""
        try:
            with open(source_path, 'r', encoding='utf-8', errors='ignore') as source_file:
                content = source_file.read()
            
            file_ext = os.path.splitext(source_path)[1]
            markdown_content = f"""# Document: {os.path.basename(source_path)}

**File Type:** {file_ext}
**Conversion Method:** Fallback text extraction

```
{content}
```
"""
            
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write(markdown_content)
            
            result = {
                "method": "fallback-text",
                "content_length": len(markdown_content),
                "original_length": len(content),
                "file_extension": file_ext
            }
            
            logger.info("Fallback conversion completed", result=result)
            return result
            
        except Exception as e:
            logger.error("Fallback conversion failed", error=str(e))
            raise DocumentConversionError(f"Fallback conversion failed: {e}")
    
    async def start(self):
        """Start the worker."""
        if not self.worker:
            raise RuntimeError("Worker not initialized. Call setup() first.")
        
        logger.info("Starting simple document converter worker...")
        await self.worker.run()
    
    async def cleanup(self):
        """Cleanup resources."""
        if self.worker:
            await self.worker.close()
        
        if self.redis_connection:
            await self.redis_connection.close()
        
        logger.info("Simple document converter worker cleaned up")


async def main():
    """Main function to run the simple document converter worker."""
    worker = SimpleDocumentConverterWorker()
    
    try:
        await worker.setup()
        await worker.start()
    except KeyboardInterrupt:
        logger.info("Received interrupt signal")
    except Exception as e:
        logger.error("Worker failed", error=str(e))
        raise
    finally:
        await worker.cleanup()


if __name__ == "__main__":
    asyncio.run(main()) 