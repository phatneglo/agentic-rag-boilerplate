"""
File Display Agent - Handles file reading, display, and content operations
"""

import asyncio
from typing import Dict, List, Any, Optional
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

from app.agents.base_agent import BaseAgent, AgentResponse, ArtifactType, AgentCapability
from app.core.config import get_settings
from app.core.logging_config import get_logger

logger = get_logger(__name__)
settings = get_settings()


class FileDisplayAgent(BaseAgent):
    """
    File Display Agent for reading and displaying file contents.
    Handles file content extraction, formatting, and presentation.
    """
    
    def __init__(self):
        # Define capabilities for this agent
        capabilities = [
            AgentCapability(
                name="File Operations",
                description="File reading, display, and content operations",
                artifact_types=[ArtifactType.CODE, ArtifactType.DOCUMENT],
                keywords=["file", "read", "display", "content", "extract", "parse", "view", "open", "document"],
                examples=["Read file contents", "Display document", "Extract text", "Parse file format"]
            )
        ]
        
        super().__init__("File Display", capabilities)
        
    def get_system_prompt(self) -> str:
        """Return the system prompt for the File Display Agent."""
        return """You are a File Display Agent. You help users with reading, displaying, and analyzing file contents.

Your capabilities include:
- Reading and displaying various file formats (PDF, DOCX, TXT, CSV, JSON, XML, etc.)
- Content extraction and text parsing
- File format conversion and processing
- Content analysis and summarization
- File structure analysis
- Metadata extraction
- Text formatting and presentation
- File content search and filtering

Provide helpful responses about file operations. When users need to read or analyze files, guide them through the process and explain file formats and content structures.

Always be practical and help users understand file contents effectively."""
        
    async def can_handle(self, user_input: str, context: Dict[str, Any] = None) -> bool:
        """Check if this agent can handle the user's request."""
        
        # Keywords that indicate file display/reading tasks
        file_display_keywords = [
            "read file", "show file", "display file", "open file", "view file",
            "file content", "file contents", "read content", "show content",
            "display content", "view content", "preview file", "extract text",
            "parse file", "process file", "analyze file content", "file viewer",
            "file reader", "content extraction", "text extraction"
        ]
        
        user_input_lower = user_input.lower()
        
        # Check for direct file display keywords
        if any(keyword in user_input_lower for keyword in file_display_keywords):
            return True
            
        # Check for file reading patterns
        reading_patterns = [
            "what's in", "what is in", "show me the", "display the",
            "read the", "open the", "view the", "extract from",
            "get content", "file analysis", "content analysis"
        ]
        
        # Check if patterns are followed by file-related words
        file_words = ["file", "document", "text", "pdf", "docx", "csv", "json", "xml"]
        
        for pattern in reading_patterns:
            if pattern in user_input_lower:
                # Check if followed by file words
                pattern_index = user_input_lower.find(pattern)
                remaining_text = user_input_lower[pattern_index + len(pattern):]
                if any(word in remaining_text for word in file_words):
                    return True
        
        return False

    async def process_request(self, user_input: str, context: Dict[str, Any] = None, config: Dict[str, Any] = None) -> AgentResponse:
        """Process file display and reading requests."""
        
        try:
            logger.info(f"File Display Agent processing: {user_input}")
            
            # Use the base class generate_response method for consistent streaming
            response_content = await self.generate_response(user_input, config=config)
            
            # Determine if we need to create file processing artifacts
            artifacts = []
            
            # Check if response involves file operations that need artifacts
            response_lower = response_content.lower()
            if any(keyword in response_lower for keyword in ["read", "extract", "parse", "process", "analyze"]):
                artifacts.append({
                    "type": ArtifactType.CODE,
                    "title": "File Processing Script",
                    "content": self._generate_file_processing_script(user_input, response_content),
                    "language": "python"
                })
            
            return AgentResponse(
                success=True,
                content=response_content,
                artifacts=artifacts,
                metadata={
                    "agent": "file_display",
                    "capabilities": ["file_reading", "content_extraction", "file_analysis", "text_processing"],
                    "supported_formats": ["pdf", "docx", "txt", "csv", "json", "xml", "html"]
                }
            )
            
        except Exception as e:
            logger.error(f"Error in File Display Agent: {e}")
            return AgentResponse(
                success=False,
                content="",
                artifacts=[],
                metadata={"agent": "file_display"},
                error=str(e)
            )
    
    def _generate_file_processing_script(self, user_input: str, response_content: str) -> str:
        """Generate a file processing script based on the request."""
        
        # Basic file processing operations template
        script_template = '''"""
File Processing Script
Generated based on: {user_input}
"""

import os
import json
import csv
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Any, Optional
import mimetypes

# PDF processing
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    print("Install: pip install PyPDF2 pdfplumber")

# DOCX processing
try:
    from docx import Document
except ImportError:
    print("Install: pip install python-docx")

# Excel processing
try:
    import pandas as pd
except ImportError:
    print("Install: pip install pandas openpyxl")

class FileProcessor:
    def __init__(self):
        self.supported_formats = {{
            '.txt': self.read_text_file,
            '.pdf': self.read_pdf_file,
            '.docx': self.read_docx_file,
            '.csv': self.read_csv_file,
            '.json': self.read_json_file,
            '.xml': self.read_xml_file,
            '.html': self.read_html_file,
            '.xlsx': self.read_excel_file,
            '.xls': self.read_excel_file
        }}
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type based on extension and MIME type."""
        file_ext = Path(file_path).suffix.lower()
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {{
            'extension': file_ext,
            'mime_type': mime_type,
            'supported': file_ext in self.supported_formats
        }}
    
    def read_file(self, file_path: str) -> Dict[str, Any]:
        """Read file content based on its type."""
        if not os.path.exists(file_path):
            return {{'error': 'File not found', 'content': None}}
        
        file_info = self.detect_file_type(file_path)
        file_ext = file_info['extension']
        
        if not file_info['supported']:
            return {{'error': f'Unsupported file type: {{file_ext}}', 'content': None}}
        
        try:
            reader_func = self.supported_formats[file_ext]
            content = reader_func(file_path)
            
            return {{
                'success': True,
                'content': content,
                'file_info': file_info,
                'size': os.path.getsize(file_path)
            }}
        except Exception as e:
            return {{'error': str(e), 'content': None}}
    
    def read_text_file(self, file_path: str) -> str:
        """Read plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def read_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """Read PDF file content."""
        try:
            # Using pdfplumber for better text extraction
            import pdfplumber
            
            text_content = []
            metadata = {{}}
            
            with pdfplumber.open(file_path) as pdf:
                metadata['num_pages'] = len(pdf.pages)
                metadata['metadata'] = pdf.metadata
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append({{
                            'page': page_num,
                            'text': page_text
                        }})
            
            return {{
                'text': '\\n\\n'.join([p['text'] for p in text_content]),
                'pages': text_content,
                'metadata': metadata
            }}
        
        except Exception as e:
            # Fallback to PyPDF2
            try:
                import PyPDF2
                
                text_content = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(reader.pages, 1):
                        text = page.extract_text()
                        text_content.append(text)
                
                return {{
                    'text': '\\n\\n'.join(text_content),
                    'num_pages': len(text_content)
                }}
            except:
                raise Exception(f"Could not read PDF: {{e}}")
    
    def read_docx_file(self, file_path: str) -> Dict[str, Any]:
        """Read DOCX file content."""
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_content.append(table_data)
        
        return {{
            'text': '\\n\\n'.join(paragraphs),
            'paragraphs': paragraphs,
            'tables': tables_content,
            'num_paragraphs': len(paragraphs),
            'num_tables': len(tables_content)
        }}
    
    def read_csv_file(self, file_path: str) -> Dict[str, Any]:
        """Read CSV file content."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            
            return {{
                'data': df.to_dict('records'),
                'columns': df.columns.tolist(),
                'shape': df.shape,
                'summary': df.describe().to_dict() if df.select_dtypes(include=[float, int]).shape[1] > 0 else None
            }}
        except:
            # Fallback to standard csv module
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                data = list(reader)
                columns = reader.fieldnames
            
            return {{
                'data': data,
                'columns': columns,
                'num_rows': len(data)
            }}
    
    def read_json_file(self, file_path: str) -> Any:
        """Read JSON file content."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def read_xml_file(self, file_path: str) -> Dict[str, Any]:
        """Read XML file content."""
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        def xml_to_dict(element):
            result = {{}}
            
            # Add attributes
            if element.attrib:
                result['@attributes'] = element.attrib
            
            # Add text content
            if element.text and element.text.strip():
                if len(element) == 0:
                    return element.text.strip()
                result['#text'] = element.text.strip()
            
            # Add child elements
            for child in element:
                child_data = xml_to_dict(child)
                if child.tag in result:
                    if not isinstance(result[child.tag], list):
                        result[child.tag] = [result[child.tag]]
                    result[child.tag].append(child_data)
                else:
                    result[child.tag] = child_data
            
            return result
        
        return {{
            'root_tag': root.tag,
            'data': xml_to_dict(root)
        }}
    
    def read_html_file(self, file_path: str) -> str:
        """Read HTML file content."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def read_excel_file(self, file_path: str) -> Dict[str, Any]:
        """Read Excel file content."""
        import pandas as pd
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheets_data = {{}}
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheets_data[sheet_name] = {{
                'data': df.to_dict('records'),
                'columns': df.columns.tolist(),
                'shape': df.shape
            }}
        
        return {{
            'sheets': sheets_data,
            'sheet_names': excel_file.sheet_names
        }}

# Example usage
def main():
    processor = FileProcessor()
    
    # Example file path
    file_path = "example.txt"  # Replace with actual file path
    
    # Check file type
    file_info = processor.detect_file_type(file_path)
    print(f"File info: {{file_info}}")
    
    # Read file content
    if file_info['supported']:
        result = processor.read_file(file_path)
        if result.get('success'):
            print(f"File size: {{result['size']}} bytes")
            print(f"Content preview: {{str(result['content'])[:200]}}...")
        else:
            print(f"Error: {{result['error']}}")
    else:
        print("Unsupported file type")

if __name__ == "__main__":
    main()
'''
        
        return script_template.format(user_input=user_input)

    async def get_capabilities(self) -> List[str]:
        """Return list of agent capabilities."""
        return [
            "file_reading",
            "content_extraction",
            "text_processing",
            "file_analysis",
            "format_conversion",
            "metadata_extraction",
            "content_formatting",
            "file_preview"
        ] 